# processors.py
"""
Обработчики текста: OCR, транскрибация, кружочки, коррекция, саммари, диалог, экспорт
Версия 5.1 — убрана обработка видеофайлов (только кружочки), добавлен DOCX
"""

import io
import os
import json
import logging
import base64
import asyncio
import subprocess
import mimetypes
import re
import time
from typing import Optional, Tuple, List, Dict, Any, AsyncGenerator
from datetime import timedelta
from openai import AsyncOpenAI

import config

# Попытка импорта дополнительных библиотек
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import docx as python_docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

# Хранилище для диалогов о документах
document_dialogues: Dict[int, Dict[int, Dict[str, Any]]] = {}


# ============================================================================
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ДЛЯ GROQ
# ============================================================================

async def _make_groq_request(groq_clients: list, func, *args, **kwargs):
    """Делаем запрос с перебором ключей и улучшенной обработкой ошибок"""
    if not groq_clients:
        raise Exception("Нет доступных Groq клиентов")

    errors = []
    client_count = len(groq_clients)

    for attempt in range(client_count * config.GROQ_RETRY_COUNT):
        client_index = attempt % client_count
        client = groq_clients[client_index]
        try:
            logger.debug(f"Попытка {attempt + 1} с клиентом {client_index}")
            return await func(client, *args, **kwargs)
        except Exception as e:
            error_msg = str(e)
            errors.append(f"Клиент {client_index}: {error_msg[:100]}")
            logger.warning(f"Ошибка запроса (попытка {attempt + 1}): {error_msg[:100]}")
            if "429" in error_msg or "rate_limit" in error_msg.lower():
                wait_time = 5 + (attempt * 2)
                logger.info(f"Rate limit, ждем {wait_time}с...")
                await asyncio.sleep(wait_time)
            else:
                await asyncio.sleep(1 + (attempt % 3))

    raise Exception(f"Все клиенты недоступны: {'; '.join(errors[:3])}")


def _truncate_text_for_model(text: str, model_type: str) -> str:
    model_limits = {
        "basic": 5000,
        "premium": 10000,
        "reasoning": 25000,
    }
    limit = model_limits.get(model_type, 5000)
    if len(text) > limit:
        logger.warning(f"Текст обрезан с {len(text)} до {limit} символов для {model_type}")
        return text[:limit] + "... [текст обрезан из-за лимитов API]"
    return text


# ============================================================================
# VISION PROCESSOR (OCR)
# ============================================================================

class VisionProcessor:
    def __init__(self):
        self.groq_clients = []

    def init_clients(self, groq_clients: list):
        self.groq_clients = groq_clients

    async def extract_text(self, image_bytes: bytes) -> str:
        if not self.groq_clients:
            return config.ERROR_NO_GROQ

        base64_image = base64.b64encode(image_bytes).decode('utf-8')

        async def extract(client):
            response = await client.chat.completions.create(
                model=config.GROQ_MODELS["vision"],
                messages=[{
                    "role": "user",
                    "content": [
                        {"type": "text", "text": config.OCR_PROMPT},
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                    ]
                }],
                temperature=config.VISION_TEMPERATURE,
                max_tokens=config.VISION_MAX_TOKENS,
            )
            return response.choices[0].message.content

        try:
            return await _make_groq_request(self.groq_clients, extract)
        except Exception as e:
            logger.error(f"Vision OCR error: {e}")
            return f"❌ Ошибка распознавания текста: {str(e)[:100]}"


vision_processor = VisionProcessor()


# ============================================================================
# VIDEO PROCESSING (только локальные файлы)
# ============================================================================

class VideoProcessor:
    @staticmethod
    async def check_video_duration(filepath: str) -> Optional[float]:
        try:
            result = subprocess.run(
                ['ffprobe', '-v', 'error', '-show_entries', 'format=duration',
                 '-of', 'default=noprint_wrappers=1:nokey=1', filepath],
                capture_output=True, text=True, timeout=10
            )
            if result.returncode == 0 and result.stdout.strip():
                return float(result.stdout.strip())
        except Exception as e:
            logger.warning(f"Error checking video duration: {e}")
        return None

    @staticmethod
    async def extract_audio_from_video(video_path: str, output_path: str) -> bool:
        try:
            subprocess.run(
                ['ffmpeg', '-i', video_path, '-vn', '-acodec', 'libmp3lame',
                 '-ab', '64k', '-ar', str(config.AUDIO_SAMPLE_RATE), '-ac', '1', '-y', output_path],
                capture_output=True, timeout=300
            )
            return os.path.exists(output_path) and os.path.getsize(output_path) > 0
        except Exception as e:
            logger.error(f"Audio extraction error: {e}")
            return False


video_processor = VideoProcessor()


# ============================================================================
# AUDIO TRANSCRIPTION
# ============================================================================

def _format_timecode(seconds: float) -> str:
    seconds = int(seconds)
    h = seconds // 3600
    m = (seconds % 3600) // 60
    s = seconds % 60
    if h > 0:
        return f"[{h:02d}:{m:02d}:{s:02d}]"
    return f"[{m:02d}:{s:02d}]"


def _segments_to_timecoded_text(segments: list) -> str:
    lines = []
    for seg in segments:
        tc = _format_timecode(seg.get("start", 0))
        text = seg.get("text", "").strip()
        if text:
            lines.append(f"{tc} {text}")
    return "\n".join(lines)


async def transcribe_voice(audio_bytes: bytes, groq_clients: list, with_timecodes: bool = False) -> str:
    async def transcribe(client):
        if with_timecodes:
            response = await client.audio.transcriptions.create(
                model=config.GROQ_MODELS["transcription"],
                file=("audio.ogg", audio_bytes, "audio/ogg"),
                language=config.AUDIO_LANGUAGE,
                response_format="verbose_json",
                temperature=config.MODEL_TEMPERATURES["transcription"],
            )
            segments = getattr(response, "segments", None)
            if segments:
                return _segments_to_timecoded_text(segments)
            return getattr(response, "text", str(response))
        else:
            response = await client.audio.transcriptions.create(
                model=config.GROQ_MODELS["transcription"],
                file=("audio.ogg", audio_bytes, "audio/ogg"),
                language=config.AUDIO_LANGUAGE,
                response_format="text",
                temperature=config.MODEL_TEMPERATURES["transcription"],
            )
            return response

    try:
        return await _make_groq_request(groq_clients, transcribe)
    except Exception as e:
        logger.error(f"Transcription error: {e}")
        return f"❌ Ошибка распознавания: {str(e)[:100]}"


# ============================================================================
# TEXT PROCESSING - CORRECTION
# ============================================================================

async def correct_text_basic(text: str, groq_clients: list) -> str:
    if not text.strip():
        return config.ERROR_EMPTY_TEXT
    text = _truncate_text_for_model(text, "basic")

    async def correct(client):
        response = await client.chat.completions.create(
            model=config.GROQ_MODELS["basic"],
            messages=[{"role": "user", "content": config.BASIC_CORRECTION_PROMPT + f"\n\nТекст:\n{text}"}],
            temperature=config.MODEL_TEMPERATURES["basic"],
        )
        return response.choices[0].message.content.strip()

    try:
        return await _make_groq_request(groq_clients, correct)
    except Exception as e:
        logger.error(f"Basic correction error: {e}")
        if "413" in str(e) or "rate_limit_exceeded" in str(e):
            shorter = text[:3000] + "... [обрезано]"
            async def retry(client):
                r = await client.chat.completions.create(
                    model=config.GROQ_MODELS["basic"],
                    messages=[{"role": "user", "content": config.BASIC_CORRECTION_PROMPT + f"\n\nТекст:\n{shorter}"}],
                    temperature=config.MODEL_TEMPERATURES["basic"],
                )
                return r.choices[0].message.content.strip()
            return await _make_groq_request(groq_clients, retry)
        return f"❌ Ошибка коррекции: {str(e)[:100]}"


async def correct_text_premium(text: str, groq_clients: list) -> str:
    if not text.strip():
        return config.ERROR_EMPTY_TEXT
    text = _truncate_text_for_model(text, "premium")

    async def correct(client):
        response = await client.chat.completions.create(
            model=config.GROQ_MODELS["premium"],
            messages=[{"role": "user", "content": config.PREMIUM_CORRECTION_PROMPT + f"\n\nТекст:\n{text}"}],
            temperature=config.MODEL_TEMPERATURES["premium"],
        )
        return response.choices[0].message.content.strip()

    try:
        return await _make_groq_request(groq_clients, correct)
    except Exception as e:
        logger.error(f"Premium correction error: {e}")
        if "413" in str(e) or "rate_limit_exceeded" in str(e):
            shorter = text[:5000] + "... [обрезано]"
            async def retry(client):
                r = await client.chat.completions.create(
                    model=config.GROQ_MODELS["premium"],
                    messages=[{"role": "user", "content": config.PREMIUM_CORRECTION_PROMPT + f"\n\nТекст:\n{shorter}"}],
                    temperature=config.MODEL_TEMPERATURES["premium"],
                )
                return r.choices[0].message.content.strip()
            return await _make_groq_request(groq_clients, retry)
        return f"❌ Ошибка коррекции: {str(e)[:100]}"


# ============================================================================
# TEXT PROCESSING - SUMMARIZATION
# ============================================================================

async def summarize_text(text: str, groq_clients: list) -> str:
    if not text.strip():
        return config.ERROR_EMPTY_TEXT

    words_count = len(text.split())
    if words_count < config.MIN_WORDS_FOR_SUMMARY or len(text) < config.MIN_CHARS_FOR_SUMMARY:
        return config.ERROR_TEXT_TOO_SHORT_FOR_SUMMARY

    text = _truncate_text_for_model(text, "reasoning")

    async def summarize(client):
        response = await client.chat.completions.create(
            model=config.GROQ_MODELS["reasoning"],
            messages=[{"role": "user", "content": config.SUMMARIZATION_PROMPT + f"\n\nТекст:\n{text}"}],
            temperature=config.MODEL_TEMPERATURES["reasoning"],
        )
        return response.choices[0].message.content.strip()

    try:
        return await _make_groq_request(groq_clients, summarize)
    except Exception as e:
        logger.error(f"Summarization error: {e}")
        if "413" in str(e) or "rate_limit_exceeded" in str(e):
            shorter = text[:10000] + "... [обрезано]"
            async def retry(client):
                r = await client.chat.completions.create(
                    model=config.GROQ_MODELS["reasoning"],
                    messages=[{"role": "user", "content": config.SUMMARIZATION_PROMPT + f"\n\nТекст:\n{shorter}"}],
                    temperature=config.MODEL_TEMPERATURES["reasoning"],
                )
                return r.choices[0].message.content.strip()
            return await _make_groq_request(groq_clients, retry)
        return f"❌ Ошибка создания саммари: {str(e)[:100]}"


# ============================================================================
# YOUTUBE СУБТИТРЫ
# ============================================================================

try:
    from youtube_transcript_api import YouTubeTranscriptApi
    YT_TRANSCRIPT_AVAILABLE = True
except ImportError:
    YT_TRANSCRIPT_AVAILABLE = False


def extract_youtube_video_id(url: str) -> Optional[str]:
    """Извлекает video_id из любого формата YouTube-ссылки."""
    patterns = [
        r'(?:youtube\.com/watch\?v=|youtu\.be/|youtube\.com/embed/|youtube\.com/shorts/)([a-zA-Z0-9_-]{11})',
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return None


def is_youtube_url(url: str) -> bool:
    return bool(extract_youtube_video_id(url.strip()))


def _format_yt_timecode(seconds: float) -> str:
    seconds = int(seconds)
    h = seconds // 3600
    m = (seconds % 3600) // 60
    s = seconds % 60
    if h > 0:
        return f"[{h:02d}:{m:02d}:{s:02d}]"
    return f"[{m:02d}:{s:02d}]"


async def fetch_youtube_subtitles(video_id: str) -> dict:
    """
    Загружает субтитры YouTube видео.
    Совместимо с youtube-transcript-api >= 1.0
    Новый API: YouTubeTranscriptApi().fetch(video_id, languages=[...])
    """
    if not YT_TRANSCRIPT_AVAILABLE:
        return {"error": "❌ Для YouTube субтитров требуется установить youtube-transcript-api"}

    def _fetch():
        ytt = YouTubeTranscriptApi()
        try:
            result = ytt.fetch(video_id, languages=["ru", "en"])
        except Exception:
            result = ytt.fetch(video_id)

        lang = getattr(result, "language_code", "unknown")

        # FetchedTranscript можно итерировать напрямую или через to_raw_data()
        try:
            segments = result.to_raw_data()
        except AttributeError:
            # Если to_raw_data нет — итерируем как список объектов
            segments = [
                {"text": s.text, "start": s.start, "duration": getattr(s, "duration", 0)}
                for s in result
            ]

        return segments, lang

    try:
        segments, lang = await asyncio.to_thread(_fetch)

        if not segments:
            return {"error": "❌ Субтитры пустые"}

        # Уточняем язык через langdetect
        sample = " ".join(s.get("text", "") for s in segments[:20])
        detected = detect_language(sample)
        if detected != "unknown":
            lang = detected

        return {"raw": segments, "lang": lang, "error": None}

    except Exception as e:
        err = str(e)
        logger.error(f"YouTube subtitles error (type={type(e).__name__}): {err}")
        if "disabled" in err.lower():
            return {"error": "❌ Субтитры отключены автором видео"}
        if "No transcripts" in err or "Could not retrieve a transcript" in err:
            return {"error": "❌ Субтитры недоступны для этого видео"}
        # Показываем реальную ошибку — поможет в отладке
        return {"error": f"❌ Ошибка субтитров: {err[:200]}"}


def _segments_to_plain_text(segments: list) -> str:
    """Субтитры → сплошной текст для передачи в LLM."""
    return " ".join(s.get("text", "").replace("\n", " ").strip() for s in segments if s.get("text", "").strip())


def _segments_to_timecoded(segments: list) -> str:
    """Субтитры → текст с таймкодами."""
    lines = []
    for seg in segments:
        text = seg.get("text", "").replace("\n", " ").strip()
        if text:
            tc = _format_yt_timecode(seg.get("start", 0))
            lines.append(f"{tc} {text}")
    return "\n".join(lines)


async def format_subtitles_as_dialogue(raw_text: str, groq_clients: list) -> str:
    """
    LLM форматирует субтитры в читаемый диалог:
    - убирает рекламные интеграции
    - группирует реплики в абзацы по смыслу
    - сохраняет живую речь
    """
    truncated = raw_text[:12000] + ("... [обрезано]" if len(raw_text) > 12000 else "")

    prompt = f"""Перед тобой субтитры видео — сплошной текст из кусочков речи.

ЗАДАЧА:
Отформатируй в читаемый текст диалога/монолога:

1. Убери рекламные интеграции — фрагменты где говорящий явно рекламирует продукт, сервис или просит подписаться/поставить лайк. Замени их на: [реклама вырезана]
2. Объедини короткие обрывки в связные абзацы по смыслу (смена темы = новый абзац)
3. Исправь только явные ошибки распознавания — не редактируй стиль и речь
4. Сохрани все имена, факты, цифры
5. Если в тексте чётко слышно смену говорящего (по контексту) — можешь обозначить абзац с новой строки, но не придумывай имена

ФОРМАТ: только готовый текст, без предисловий

Субтитры:
{truncated}"""

    async def fmt(client):
        response = await client.chat.completions.create(
            model=config.GROQ_MODELS["premium"],
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
        )
        return response.choices[0].message.content.strip()

    try:
        return await _make_groq_request(groq_clients, fmt)
    except Exception as e:
        logger.error(f"Subtitle formatting error: {e}")
        # Fallback — возвращаем сырой текст
        return raw_text


# ============================================================================
# URL SCRAPING
# ============================================================================

def is_url(text: str) -> bool:
    """Проверяет, является ли текст обычным URL (не YouTube)."""
    text = text.strip()
    if not (text.startswith(("http://", "https://")) and " " not in text and len(text) > 10):
        return False
    return not is_youtube_url(text)


async def fetch_url_text(url: str) -> str:
    """Скачивает страницу по URL и извлекает текст. С retry при 429."""
    try:
        import httpx
        from html.parser import HTMLParser

        class _TextExtractor(HTMLParser):
            def __init__(self):
                super().__init__()
                self._text = []
                self._skip = False

            def handle_starttag(self, tag, attrs):
                if tag in ("script", "style", "nav", "footer", "header", "aside", "menu"):
                    self._skip = True

            def handle_endtag(self, tag):
                if tag in ("script", "style", "nav", "footer", "header", "aside", "menu"):
                    self._skip = False

            def handle_data(self, data):
                if not self._skip:
                    stripped = data.strip()
                    if stripped:
                        self._text.append(stripped)

            def get_text(self):
                return "\n".join(self._text)

        # Несколько вариантов User-Agent для ротации
        user_agents = [
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36",
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36",
            "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36",
        ]

        import random
        headers = {
            "User-Agent": random.choice(user_agents),
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8",
            "Accept-Language": "ru-RU,ru;q=0.9,en-US;q=0.8,en;q=0.7",
            "Accept-Encoding": "gzip, deflate, br",
            "DNT": "1",
            "Upgrade-Insecure-Requests": "1",
            "Sec-Fetch-Dest": "document",
            "Sec-Fetch-Mode": "navigate",
            "Sec-Fetch-Site": "none",
            "Cache-Control": "max-age=0",
        }

        last_error = None
        for attempt in range(3):
            if attempt > 0:
                await asyncio.sleep(2 * attempt)  # 2s, 4s между попытками
                headers["User-Agent"] = random.choice(user_agents)

            try:
                async with httpx.AsyncClient(
                    timeout=20,
                    follow_redirects=True,
                    headers=headers,
                ) as client:
                    response = await client.get(url)

                    if response.status_code == 429:
                        retry_after = int(response.headers.get("Retry-After", 5))
                        wait = min(retry_after, 10)
                        logger.warning(f"URL 429, waiting {wait}s (attempt {attempt+1})")
                        await asyncio.sleep(wait)
                        last_error = f"429 Too Many Requests"
                        continue

                    if response.status_code == 403:
                        return "❌ Сайт закрыт для автоматических запросов (403 Forbidden)"

                    if response.status_code == 401:
                        return "❌ Сайт требует авторизации"

                    response.raise_for_status()

                    parser = _TextExtractor()
                    parser.feed(response.text)
                    text = parser.get_text()

                    lines = [l.strip() for l in text.splitlines() if len(l.strip()) > 20]
                    text = "\n".join(lines)

                    if not text or len(text) < 100:
                        return "❌ Не удалось извлечь текст со страницы. Возможно, контент загружается динамически (JavaScript)."

                    if len(text) > 30000:
                        text = text[:30000] + "\n... [страница обрезана]"

                    logger.info(f"Fetched URL {url}: {len(text)} chars")
                    return text

            except httpx.TimeoutException:
                last_error = "таймаут соединения"
                continue
            except httpx.HTTPStatusError as e:
                last_error = str(e)
                break

        return f"❌ Не удалось загрузить страницу: {last_error or 'неизвестная ошибка'}"

    except ImportError:
        return "❌ Для обработки ссылок требуется установить httpx"
    except Exception as e:
        logger.error(f"URL fetch error: {e}")
        return f"❌ Не удалось загрузить страницу: {str(e)[:100]}"


# ============================================================================
# ОПРЕДЕЛЕНИЕ ЯЗЫКА
# ============================================================================

def detect_language(text: str) -> str:
    """Определяет язык текста. Возвращает код ('ru', 'en', ...) или 'unknown'."""
    try:
        from langdetect import detect
        return detect(text[:1000])
    except Exception:
        return "unknown"


def is_non_russian(text: str) -> bool:
    """Возвращает True если текст явно не на русском."""
    return detect_language(text) not in ("ru", "unknown")


# ============================================================================
# ПЕРЕВОД
# ============================================================================

async def translate_to_russian(text: str, groq_clients: list) -> str:
    """Переводит текст на русский язык."""
    if not text.strip():
        return config.ERROR_EMPTY_TEXT

    text_to_translate = _truncate_text_for_model(text, "premium")

    prompt = (
        "Переведи следующий текст на русский язык.\n"
        "Сохрани структуру, абзацы и форматирование оригинала.\n"
        "Переводи точно, без сокращений и добавлений.\n"
        "Выведи ТОЛЬКО перевод, без предисловий и комментариев.\n\n"
        f"Текст:\n{text_to_translate}"
    )

    async def translate(client):
        response = await client.chat.completions.create(
            model=config.GROQ_MODELS["premium"],
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
        )
        return response.choices[0].message.content.strip()

    try:
        return await _make_groq_request(groq_clients, translate)
    except Exception as e:
        logger.error(f"Translation error: {e}")
        return f"❌ Ошибка перевода: {str(e)[:100]}"


# ============================================================================
# РАБОТА НАД ОШИБКАМИ
# ============================================================================

async def explain_corrections(original_text: str, corrected_text: str, groq_clients: list) -> str:
    """
    Сравнивает оригинал и исправленный текст, объясняет каждую правку.
    Использует premium модель — она точнее находит различия.
    """
    if not original_text.strip() or not corrected_text.strip():
        return "❌ Нет текста для разбора"

    # Если тексты идентичны — сразу говорим об этом
    if original_text.strip() == corrected_text.strip():
        return "🎓 <b>Работа над ошибками</b>\n\nОшибок не найдено — текст был чистым ✅"

    # Обрезаем оба текста чтобы уложиться в лимит модели
    max_len = 4000
    orig_truncated = original_text[:max_len] + ("... [обрезан]" if len(original_text) > max_len else "")
    corr_truncated = corrected_text[:max_len] + ("... [обрезан]" if len(corrected_text) > max_len else "")

    prompt = (
        config.EXPLAIN_CORRECTIONS_PROMPT
        + f"\n\nОРИГИНАЛ:\n{orig_truncated}"
        + f"\n\nИСПРАВЛЕННЫЙ ТЕКСТ:\n{corr_truncated}"
    )

    async def explain(client):
        response = await client.chat.completions.create(
            model=config.GROQ_MODELS["premium"],
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,   # низкая температура — нужна точность, не творчество
            max_tokens=2000,
        )
        return response.choices[0].message.content.strip()

    try:
        return await _make_groq_request(groq_clients, explain)
    except Exception as e:
        logger.error(f"Explain corrections error: {e}")
        return f"❌ Ошибка при разборе правок: {str(e)[:100]}"


# ============================================================================
# ДИАЛОГОВЫЙ РЕЖИМ
# ============================================================================

def save_document_for_dialog(user_id: int, msg_id: int, document_text: str, source: str = "unknown"):
    if user_id not in document_dialogues:
        document_dialogues[user_id] = {}
    document_dialogues[user_id][msg_id] = {
        "full_text": document_text,
        "text": document_text,
        "original": document_text,
        "history": [],
        "timestamp": time.time(),
        "source": source
    }
    logger.info(f"💾 Документ для диалога: user={user_id}, msg={msg_id}, len={len(document_text)}")
    return document_dialogues[user_id][msg_id]


def get_document_text(user_id: int, msg_id: int) -> Optional[str]:
    if user_id not in document_dialogues or msg_id not in document_dialogues[user_id]:
        return None
    doc_data = document_dialogues[user_id][msg_id]
    for key in ["full_text", "text", "original"]:
        if key in doc_data and doc_data[key]:
            return doc_data[key]
    return None


async def stream_document_answer(
    user_id: int,
    msg_id: int,
    question: str,
    groq_clients: list
) -> AsyncGenerator[str, None]:
    if not groq_clients:
        yield "❌ Нет доступных Groq клиентов"
        return

    if user_id not in document_dialogues or msg_id not in document_dialogues[user_id]:
        yield "❌ Документ не найден. Сначала загрузите документ."
        return

    doc_data = document_dialogues[user_id][msg_id]
    full_text = get_document_text(user_id, msg_id)
    if not full_text:
        yield "❌ Не удалось извлечь текст документа."
        return

    history = doc_data.get("history", [])
    context = ""
    for turn in history[-5:]:
        q = turn.get('question') or turn.get('q', '')
        a = turn.get('answer') or turn.get('a', '')
        if q and a:
            context += f"Вопрос: {q}\nОтвет: {a}\n\n"

    doc_preview = full_text[:20000] + "... [обрезан]" if len(full_text) > 20000 else full_text

    prompt = f"""Ты — ассистент, который отвечает на вопросы по содержанию документа.

Документ:
{doc_preview}

{context}

Вопрос:
{question}

Ответь на вопрос, используя только информацию из документа. Если ответа нет в документе, так и скажи.
Ответ должен быть подробным, но по существу."""

    client = groq_clients[0 % len(groq_clients)]

    try:
        stream = await client.chat.completions.create(
            model=config.GROQ_MODELS["reasoning"],
            messages=[
                {"role": "system", "content": "Ты отвечаешь строго по документу."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.2,
            stream=True,
        )

        full_answer = ""
        async for chunk in stream:
            if chunk.choices and chunk.choices[0].delta.content:
                piece = chunk.choices[0].delta.content
                full_answer += piece
                yield piece

        history.append({
            "question": question, "answer": full_answer,
            "q": question, "a": full_answer,
            "timestamp": time.time()
        })
        doc_data["history"] = history[-config.MAX_DIALOG_HISTORY:]

    except Exception as e:
        logger.error(f"Stream error: {e}", exc_info=True)
        yield f"❌ Ошибка при генерации ответа: {str(e)[:100]}"


# ============================================================================
# FILE PROCESSING
# ============================================================================

async def process_video_file(video_bytes: bytes, filename: str, groq_clients: list, with_timecodes: bool = False) -> str:
    try:
        file_ext = filename.split('.')[-1] if '.' in filename else 'mp4'
        temp_video_path = f"{config.TEMP_DIR}/video_{int(time.time())}_{os.getpid()}.{file_ext}"
        temp_audio_path = f"{config.TEMP_DIR}/audio_{int(time.time())}_{os.getpid()}.mp3"

        with open(temp_video_path, 'wb') as f:
            f.write(video_bytes)

        duration = await video_processor.check_video_duration(temp_video_path)
        if duration and duration > 3600:
            os.remove(temp_video_path)
            return config.ERROR_VIDEO_TOO_LONG

        if not await video_processor.extract_audio_from_video(temp_video_path, temp_audio_path):
            os.remove(temp_video_path)
            return "❌ Ошибка извлечения звука из видео"

        with open(temp_audio_path, 'rb') as f:
            audio_bytes = f.read()

        text = await transcribe_voice(audio_bytes, groq_clients, with_timecodes=with_timecodes)

        for p in [temp_video_path, temp_audio_path]:
            try:
                os.remove(p)
            except:
                pass

        return text

    except Exception as e:
        logger.error(f"Error processing video file: {e}")
        return f"❌ Ошибка обработки видеофайла: {str(e)[:100]}"


async def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """Извлечение текста из PDF. Тяжёлая работа вынесена в thread чтобы не блокировать event loop."""
    if not PDFPLUMBER_AVAILABLE:
        return "❌ Для работы с PDF требуется установить pdfplumber"

    def _extract_sync():
        pdf_buffer = io.BytesIO(pdf_bytes)
        text = ""
        page_count = 0

        with pdfplumber.open(pdf_buffer) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                if config.PDF_MAX_PAGES and page_num > config.PDF_MAX_PAGES:
                    break

                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- Страница {page_num} ---\n"
                    text += page_text + "\n"

                tables = page.find_tables()
                if tables:
                    for table_idx, table in enumerate(tables, 1):
                        text += f"\n[Таблица {table_idx} на странице {page_num}]\n"
                        table_data = table.extract()
                        for row in table_data:
                            if row:
                                text += " | ".join(str(cell) if cell else "" for cell in row) + "\n"

                page_count += 1

        if not text.strip():
            raise ValueError("Не удалось извлечь текст из PDF")

        logger.info(f"Extracted text from {page_count} PDF pages, {len(pdf_bytes) // 1024} KB")
        return text.strip()

    try:
        return await asyncio.to_thread(_extract_sync)
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return f"❌ Ошибка обработки PDF: {str(e)}"


async def extract_text_from_docx(docx_bytes: bytes) -> str:
    if not DOCX_AVAILABLE:
        return "❌ Для работы с DOCX требуется установить python-docx"
    try:
        doc_buffer = io.BytesIO(docx_bytes)
        doc = python_docx.Document(doc_buffer)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if not text.strip():
            return "❌ Документ пуст"
        return text.strip()
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return f"❌ Ошибка обработки DOCX: {str(e)}"


async def extract_text_from_txt(txt_bytes: bytes) -> str:
    try:
        for encoding in ['utf-8', 'cp1251', 'koi8-r', 'windows-1251']:
            try:
                return txt_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue
        return txt_bytes.decode('utf-8', errors='ignore')
    except Exception as e:
        logger.error(f"TXT reading error: {e}")
        return f"❌ Ошибка чтения текстового файла: {str(e)}"


async def extract_text_from_file(file_bytes: bytes, filename: str, groq_clients: list) -> str:
    mime_type, _ = mimetypes.guess_type(filename)
    file_ext = filename.lower().split('.')[-1] if '.' in filename else ''

    if mime_type and mime_type.startswith('image/') or file_ext in ['jpg', 'jpeg', 'png', 'bmp', 'gif', 'webp']:
        vision_processor.init_clients(groq_clients)
        return await vision_processor.extract_text(file_bytes)

    if mime_type == 'application/pdf' or file_ext == 'pdf':
        return await extract_text_from_pdf(file_bytes)

    if mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or file_ext == 'docx':
        return await extract_text_from_docx(file_bytes)

    if mime_type == 'text/plain' or file_ext == 'txt':
        return await extract_text_from_txt(file_bytes)

    if file_ext == 'doc':
        return config.ERROR_DOC_NOT_SUPPORTED

    return config.ERROR_UNSUPPORTED_FORMAT


# ============================================================================
# ЭКСПОРТ В ФАЙЛЫ
# ============================================================================

async def save_to_txt(text: str, filepath: str) -> bool:
    try:
        def _write():
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(text)
        await asyncio.to_thread(_write)
        return True
    except Exception as e:
        logger.error(f"TXT save error: {e}")
        return False


async def save_to_pdf(text: str, filepath: str) -> bool:
    try:
        def _write():
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
            from reportlab.lib.utils import simpleSplit

            c = canvas.Canvas(filepath, pagesize=A4)
            width, height = A4
            margin = 50
            line_height = 14
            y = height - margin

            c.setFont("Helvetica-Bold", 14)
            c.drawString(margin, y, "Обработанный текст")
            y -= 30
            c.setFont("Helvetica", 10)
            from datetime import datetime
            c.drawString(margin, y, f"Создано: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
            y -= 40
            c.setFont("Helvetica", 11)
            max_width = width - 2 * margin

            for paragraph in text.split('\n'):
                if not paragraph.strip():
                    y -= line_height
                    continue
                for line in simpleSplit(paragraph, "Helvetica", 11, max_width):
                    if y < margin + 20:
                        c.showPage()
                        y = height - margin
                        c.setFont("Helvetica", 11)
                    c.drawString(margin, y, line)
                    y -= line_height
            c.save()

        await asyncio.to_thread(_write)
        return True
    except ImportError:
        logger.warning("reportlab not installed, falling back to txt")
        return False
    except Exception as e:
        logger.error(f"PDF save error: {e}")
        return False


async def save_to_docx(text: str, filepath: str) -> bool:
    """Сохраняет текст в DOCX через python-docx."""
    if not DOCX_AVAILABLE:
        logger.warning("python-docx not installed")
        return False

    try:
        def _write():
            from docx import Document as DocxDocument
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from datetime import datetime

            doc = DocxDocument()

            # Заголовок
            title = doc.add_heading("Обработанный текст", level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Дата
            date_para = doc.add_paragraph(f"Создано: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
            date_para.runs[0].font.size = Pt(9)
            date_para.runs[0].font.color.rgb = None  # серый через style

            doc.add_paragraph()  # отступ

            # Основной текст: каждый абзац — отдельный параграф
            for line in text.split('\n'):
                p = doc.add_paragraph(line)
                p.runs[0].font.size = Pt(11) if p.runs else None

            doc.save(filepath)

        await asyncio.to_thread(_write)
        return True
    except Exception as e:
        logger.error(f"DOCX save error: {e}")
        return False


# ============================================================================
# РАЗБОР ПО КОСТОЧКАМ
# ============================================================================

async def breakdown_corrections(original_text: str, corrected_text: str, groq_clients: list) -> str:
    """
    Объясняет каждое исправление между оригиналом и исправленным текстом.
    Использует premium-модель для точности.
    """
    if not original_text.strip() or not corrected_text.strip():
        return "❌ Нет текста для разбора."

    if original_text.strip() == corrected_text.strip():
        return "✅ Текст был чистым — исправлений нет."

    # Обрезаем чтобы уложиться в лимит — берём оба текста
    max_chars = 4000
    orig = original_text[:max_chars]
    corr = corrected_text[:max_chars]

    prompt = (
        config.BREAKDOWN_PROMPT
        + f"\n\nОРИГИНАЛ:\n{orig}"
        + f"\n\nИСПРАВЛЕННЫЙ ТЕКСТ:\n{corr}"
    )

    async def analyze(client):
        response = await client.chat.completions.create(
            model=config.GROQ_MODELS["premium"],
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
        )
        return response.choices[0].message.content.strip()

    try:
        return await _make_groq_request(groq_clients, analyze)
    except Exception as e:
        logger.error(f"Breakdown error: {e}")
        return f"❌ Ошибка при разборе: {str(e)[:100]}"


# ============================================================================
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
# ============================================================================

def get_available_modes(text: str) -> list:
    words_count = len(text.split())
    text_length = len(text)
    available = ["basic", "premium"]
    if words_count >= config.MIN_WORDS_FOR_SUMMARY and text_length >= config.MIN_CHARS_FOR_SUMMARY:
        available.append("summary")
    return available


# ============================================================================
# ЭКСПОРТ
# ============================================================================

__all__ = [
    'transcribe_voice',
    'correct_text_basic',
    'correct_text_premium',
    'summarize_text',
    'breakdown_corrections',
    'extract_text_from_file',
    'get_available_modes',
    'vision_processor',
    'save_document_for_dialog',
    'stream_document_answer',
    'get_document_text',
    'document_dialogues',
    'save_to_txt',
    'save_to_pdf',
    'save_to_docx',
    'explain_corrections',
    'breakdown_corrections',
    'is_url',
    'fetch_url_text',
    'is_youtube_url',
    'extract_youtube_video_id',
    'fetch_youtube_subtitles',
    'format_subtitles_as_dialogue',
    'YT_TRANSCRIPT_AVAILABLE',
    'detect_language',
    'is_non_russian',
    'translate_to_russian',
    'PDFPLUMBER_AVAILABLE',
    'DOCX_AVAILABLE',
]
